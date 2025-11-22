
import os
import re
import warnings
from datetime import datetime
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from scipy import stats
from scipy.stats import wilcoxon, shapiro
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")

# ========== CONFIG ==========
INPUT_FILE = "Datos Encuesta para Personal Administrativo copy.csv"
OUTPUT_FOLDER = "resultados_tesis_unasam_2025_CORREGIDOS_final_v1"
GRAF_DIR = os.path.join(OUTPUT_FOLDER, "graficos")
os.makedirs(GRAF_DIR, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Parámetros mínimos para pruebas
MIN_N_ALFA = 3    
MIN_N_SHAPIRO = 3
MIN_N_WILCOXON = 5

print("\nIniciando análisis — revisa OUTPUT en:", OUTPUT_FOLDER)

# =========Función para extraer numero pregunta  ==========
def extraer_num_pregunta(col_name):
    """
    Extrae el primer número entero 1-99 que aparezca en el nombre de la columna.
    Devuelve int o None si no encuentra.
    Acepta formatos: "1. Pregunta", "Pregunta 01 - ...", "P1", "1) ...", etc.
    """
    if not isinstance(col_name, str):
        return None
    m = re.search(r'\b([1-9][0-9]?)\b', col_name)
    if m:
        return int(m.group(1))
    m2 = re.search(r'[pP]\s*[:\-]?\s*([1-9][0-9]?)', col_name)
    return int(m2.group(1)) if m2 else None

# ========== MAPEOS LIKERT ==========
LIKERT_MAP = {
    "totalmente en desacuerdo": 1,
    "muy en desacuerdo": 1,
    "en desacuerdo": 2,
    "1": 1, "2": 2, "3": 3, "4": 4, "5": 5,
    "ni de acuerdo / ni en desacuerdo": 3,
    "ni de acuerdo/ni en desacuerdo": 3,
    "ni de acuerdo ni en desacuerdo": 3,
    "neutral": 3,
    "de acuerdo": 4,
    "totalmente de acuerdo": 5,
    "muy de acuerdo": 5
}

def convertir_likert(val):
    """Convierte textos y números a escala 1-5, devuelve np.nan si no es convertible."""
    if pd.isna(val):
        return np.nan
    s = str(val).strip().lower()
    for key, num in LIKERT_MAP.items():
        if key == s:
            return num
    for key, num in LIKERT_MAP.items():
        if key in s and len(key) > 1:
            return num
    # extraer dígito
    m = re.search(r'([1-5])', s)
    if m:
        return int(m.group(1))
    return np.nan

# ========== OPERACIONALIZACIÓN ==========
OPERACIONALIZACION = {
    "VD": {
        "Eficiencia operativa actual": {"items": [1, 2, 3]},
        "Calidad de registros": {"items": [4, 5, 6]},
        "Carga operativa del personal": {"items": [7, 8, 9]},
        "Integración sistémica": {"items": [10, 11, 12]},
        "Seguridad y auditoría": {"items": [13, 14, 15]}
    },
    "VI": {
        "Impacto esperado en eficiencia": {"items": [16, 17, 18]},
        "Facilitación de integración sistémica": {"items": [19]},
        "Viabilidad y necesidad percibida": {"items": [20]}
    }
}

# ========== 1) CARGA Y LIMPIEZA ==========
print("\n[1] Cargando datos...")
df_raw = pd.read_csv(INPUT_FILE, encoding='utf-8', dtype=str)   
df_raw = df_raw.dropna(how="all").drop_duplicates()
n_raw = len(df_raw)
print(f"  • Filas leídas (bruto): {n_raw}")
# eliminar columnas innecesarias 
cols_to_drop_candidates = [
    "Marca temporal", "Marca de tiempo", "Timestamp", "Nombre y Apellidos", "Nombre",
    "Cargo (Opcional)", "Cargo", "Email", "Correro", "Correo"
]
cols_found_drop = [c for c in df_raw.columns if any(x.lower() in c.lower() for x in cols_to_drop_candidates)]
if cols_found_drop:
    df_raw = df_raw.drop(columns=cols_found_drop)
    print(f"  • Columnas administrativas eliminadas: {cols_found_drop}")

# limpiar nombres columnas: strip
df_raw.columns = [c.strip() for c in df_raw.columns]

# convertir a numérico Likert
df_num = df_raw.copy()
for col in df_num.columns:
    df_num[col] = df_num[col].apply(convertir_likert).astype(float)

# crear mapa columna->número de pregunta
col_to_num = {col: extraer_num_pregunta(col) for col in df_num.columns}
mapped_count = sum(1 for v in col_to_num.values() if v is not None)
print(f"  • Columnas con número de pregunta detectado: {mapped_count}/{len(col_to_num)}")

# construir dict num->col para evitar ambigüedades
num_to_col = {}
for col, num in col_to_num.items():
    if num is not None:
        # si conflicto (varias columnas con mismo número), prefiero la primera encontrada
        if num not in num_to_col:
            num_to_col[num] = col

missing_items = []
for var in OPERACIONALIZACION:
    for dim, info in OPERACIONALIZACION[var].items():
        for qnum in info["items"]:
            if qnum not in num_to_col:
                missing_items.append(qnum)
if missing_items:
    print(" Atención: faltan preguntas en el CSV para los ítems:", sorted(set(missing_items)))
    print("  - Si no usas las nuevas preguntas, el análisis no corresponderá exactamente a lo planificado.")
else:
    print("  • Todas las preguntas de la operacionalización están presentes en el CSV.")

# ========== 2) FUNCIONES DE CONFIABILIDAD ==========
def cronbach_alpha(df_items):
    """Alpha de Cronbach clásico (Pearson) sobre dataframe de ítems numericos."""
    df_items = df_items.dropna()
    k = df_items.shape[1]
    if k < 2 or df_items.shape[0] < 2:
        return np.nan
    item_vars = df_items.var(axis=0, ddof=1).sum()
    total_var = df_items.sum(axis=1).var(ddof=1)
    if total_var == 0:
        return np.nan
    alpha = (k / (k - 1)) * (1 - item_vars / total_var)
    return alpha

def alpha_ordinal(df_items):
    """
    Estimación del 'alfa ordinal' aproximada usando correlaciones de Spearman
    y la fórmula basada en la correlación inter-ítems promedio.
    """
    df_i = df_items.dropna()
    k = df_i.shape[1]
    if k < 2 or df_i.shape[0] < 2:
        return np.nan
    corr = df_i.corr(method='spearman')
    # promedio de correlaciones off-diagonal
    sum_off = corr.values.sum() - np.trace(corr.values)
    denom = k * (k - 1)
    if denom == 0:
        return np.nan
    r_bar = sum_off / denom
    # evitar división por cero
    if (1 + (k - 1) * r_bar) == 0:
        return np.nan
    alpha = (k * r_bar) / (1 + (k - 1) * r_bar)
    return alpha

# ========== 3) CALCULAR ALFAS POR DIMENSIÓN ==========
print("\n[2] Confiabilidad por dimensión (Pearson + Ordinal) ...")
rows_alpha = []
for var in OPERACIONALIZACION:
    for dim, info in OPERACIONALIZACION[var].items():
        # columnas reales para esos números de ítem
        cols = [num_to_col[q] for q in info["items"] if q in num_to_col]
        df_items = df_num[cols] if cols else pd.DataFrame()
        if df_items.shape[1] < 2 or df_items.dropna().shape[0] < MIN_N_ALFA:
            # suficiente info? si no, registrar NaNs
            rows_alpha.append({
                "Variable": "VD" if var == "VD" else "VI",
                "Dimensión": dim,
                "Ítems": info["items"],
                "N_items_presentes": df_items.shape[1],
                "N_respuestas_validas": df_items.dropna().shape[0],
                "Alfa_Pearson": np.nan,
                "Alfa_Ordinal": np.nan
            })
            continue
        a_p = cronbach_alpha(df_items)
        a_o = alpha_ordinal(df_items)
        rows_alpha.append({
            "Variable": "VD" if var == "VD" else "VI",
            "Dimensión": dim,
            "Ítems": info["items"],
            "N_items_presentes": df_items.shape[1],
            "N_respuestas_validas": df_items.dropna().shape[0],
            "Alfa_Pearson": round(a_p, 3) if not np.isnan(a_p) else np.nan,
            "Alfa_Ordinal": round(a_o, 3) if not np.isnan(a_o) else np.nan
        })

df_alpha = pd.DataFrame(rows_alpha)
df_alpha.to_csv(os.path.join(OUTPUT_FOLDER, "01_confiabilidad_cronbach.csv"), index=False, encoding='utf-8-sig')
print("  • CSV generado: 01_confiabilidad_cronbach.csv")

# ======= Reporte adicional: Alfa por Variable (VD / VI) =======
rows_alpha_var = []
for var in OPERACIONALIZACION:
    # reunir todos los números de pregunta de la variable
    qnums = []
    for dim, info in OPERACIONALIZACION[var].items():
        qnums.extend(info.get("items", []))
    cols = [num_to_col[q] for q in sorted(set(qnums)) if q in num_to_col]
    df_items = df_num[cols] if cols else pd.DataFrame()
    n_items = df_items.shape[1]
    n_valid = df_items.dropna().shape[0]
    if n_items < 2 or n_valid < MIN_N_ALFA:
        a_p = np.nan
        a_o = np.nan
    else:
        a_p = cronbach_alpha(df_items)
        a_o = alpha_ordinal(df_items)
    rows_alpha_var.append({
        "Variable": var,
        "Ítems_considerados": sorted(set(qnums)),
        "N_items_presentes": int(n_items),
        "N_respuestas_validas": int(n_valid),
        "Alfa_Pearson": round(a_p, 3) if not np.isnan(a_p) else np.nan,
        "Alfa_Ordinal": round(a_o, 3) if not np.isnan(a_o) else np.nan
    })

df_alpha_var = pd.DataFrame(rows_alpha_var)
df_alpha_var.to_csv(os.path.join(OUTPUT_FOLDER, "01_confiabilidad_por_variable.csv"), index=False, encoding='utf-8-sig')
print("  • CSV generado: 01_confiabilidad_por_variable.csv")

# ======= Reporte adicional: Alfa general de toda la encuesta (todos los ítems numerados) =======
all_nums = sorted([n for n in num_to_col.keys()])
all_cols = [num_to_col[n] for n in all_nums] if all_nums else []
df_all = df_num[all_cols] if all_cols else pd.DataFrame()
n_items_all = df_all.shape[1]
n_valid_all = df_all.dropna().shape[0]
if n_items_all < 2 or n_valid_all < MIN_N_ALFA:
    alfa_total_p = np.nan
    alfa_total_o = np.nan
else:
    alfa_total_p = cronbach_alpha(df_all)
    alfa_total_o = alpha_ordinal(df_all)

df_alfa_total = pd.DataFrame([{
    "N_items_presentes": int(n_items_all),
    "N_respuestas_validas": int(n_valid_all),
    "Alfa_Pearson_total": round(alfa_total_p, 3) if not np.isnan(alfa_total_p) else np.nan,
    "Alfa_Ordinal_total": round(alfa_total_o, 3) if not np.isnan(alfa_total_o) else np.nan
}])

df_alfa_total.to_csv(os.path.join(OUTPUT_FOLDER, "01_confiabilidad_total_encuesta.csv"), index=False, encoding='utf-8-sig')
print("  • CSV generado: 01_confiabilidad_total_encuesta.csv")

# ========== 4) Estadísticas descriptivas por ítem y por dimensión ==========
print("\n[3] Estadísticas descriptivas (ítem y dimensión)...")
rows_desc_dim = []
rows_desc_item = []

# procesar ítem por ítem (frecuencias)
for col in df_num.columns:
    serie = df_num[col].dropna().astype(int)
    if serie.empty:
        continue
    frec = serie.value_counts().reindex([1,2,3,4,5], fill_value=0)
    pct = (frec / frec.sum() * 100).round(2)
    rows_desc_item.append({
        "Columna": col,
        "Pregunta_num": col_to_num.get(col, None),
        "N": int(frec.sum()),
        "media": float(serie.mean()),
        "mediana": float(serie.median()),
        "std": float(serie.std(ddof=1)),
        "%1": float(pct.loc[1]),
        "%2": float(pct.loc[2]),
        "%3": float(pct.loc[3]),
        "%4": float(pct.loc[4]),
        "%5": float(pct.loc[5])
    })

df_items_desc = pd.DataFrame(rows_desc_item).sort_values(by="Pregunta_num")
df_items_desc.to_csv(os.path.join(OUTPUT_FOLDER, "02_desc_item_frecuencias.csv"), index=False, encoding='utf-8-sig')
print("  • CSV generado: 02_desc_item_frecuencias.csv")

# por dimensión (medianas por persona dentro de la dimensión)
for var in OPERACIONALIZACION:
    for dim, info in OPERACIONALIZACION[var].items():
        cols = [num_to_col[q] for q in info["items"] if q in num_to_col]
        if not cols:
            continue
        df_dim = df_num[cols]
        medianas_personas = df_dim.median(axis=1).dropna()
        if medianas_personas.empty:
            continue
        pct_acuerdo = (np.sum(df_dim.values.flatten() >= 4) / np.count_nonzero(~np.isnan(df_dim.values.flatten()))) * 100
        rows_desc_dim.append({
            "Variable": "VD" if var == "VD" else "VI",
            "Dimensión": dim,
            "Ítems": info["items"],
            "N_personas": int(medianas_personas.shape[0]),
            "Mediana": float(medianas_personas.median()),
            "Q1": float(medianas_personas.quantile(0.25)),
            "Q3": float(medianas_personas.quantile(0.75)),
            "IQR": float(medianas_personas.quantile(0.75) - medianas_personas.quantile(0.25)),
            "%Acuerdo(4-5)": round(pct_acuerdo, 2)
        })

df_desc_dim = pd.DataFrame(rows_desc_dim).sort_values(by="Mediana", ascending=False)
df_desc_dim.to_csv(os.path.join(OUTPUT_FOLDER, "03_desc_dimensiones.csv"), index=False, encoding='utf-8-sig')
print("  • CSV generado: 03_desc_dimensiones.csv")

# ========== 5) GRAFICOS (por ítem y por dimensión) ==========
print("\n[4] Generando gráficos individuales por ítem y por dimensión...")
# Gráficos por ítem: barra de frecuencias y porcentaje acumulado
for _, row in df_items_desc.iterrows():
    col = row["Columna"]
    serie = df_num[col].dropna().astype(int)
    if serie.empty:
        continue
    fig, ax = plt.subplots(figsize=(6,4))
    frec = serie.value_counts().reindex([1,2,3,4,5], fill_value=0)
    ax.bar([1,2,3,4,5], frec.values)
    ax.set_xlabel("Respuesta (Likert 1-5)")
    ax.set_ylabel("Frecuencia")
    ax.set_title(f"Pregunta {row['Pregunta_num']}: Distribución de respuestas")
    ax.set_xticks([1,2,3,4,5])
    plt.tight_layout()
    fname = os.path.join(GRAF_DIR, f"item_{row['Pregunta_num']}_frecuencia.png") if row['Pregunta_num'] else os.path.join(GRAF_DIR, f"item_{re.sub('[^0-9]','',col)}_frecuencia.png")
    fig.savefig(fname, dpi=200)
    plt.close(fig)

# Boxplot y violín por dimensión
for var in OPERACIONALIZACION:
    for dim, info in OPERACIONALIZACION[var].items():
        cols = [num_to_col[q] for q in info["items"] if q in num_to_col]
        if len(cols) < 1:
            continue
        df_dim = df_num[cols].dropna()
        if df_dim.empty:
            continue
        # Boxplot (por persona)
        fig, ax = plt.subplots(figsize=(6,4))
        ax.boxplot(df_dim.median(axis=1).dropna(), vert=False, labels=[dim])
        ax.set_xlabel("Mediana por persona (escala 1-5)")
        ax.set_title(f"Boxplot - {dim}")
        plt.tight_layout()
        fname_box = os.path.join(GRAF_DIR, f"dim_{re.sub('[^A-Za-z0-9]','_',dim) }_boxplot.png")
        fig.savefig(fname_box, dpi=200)
        plt.close(fig)

        # Violin por ítem dentro de la dimensión
        fig2, ax2 = plt.subplots(figsize=(8,4))
        data_for_violin = [df_dim[c].dropna().values for c in df_dim.columns]
        ax2.violinplot(data_for_violin, showmeans=True)
        ax2.set_xticks(np.arange(1, len(df_dim.columns)+1))
        ax2.set_xticklabels([f"{extraer_num_pregunta(c)}" for c in df_dim.columns])
        ax2.set_ylabel("Respuestas (1-5)")
        ax2.set_title(f"Violin por ítem - {dim} (ítems: {info['items']})")
        plt.tight_layout()
        fname_v = os.path.join(GRAF_DIR, f"dim_{re.sub('[^A-Za-z0-9]','_',dim)}_violin.png")
        fig2.savefig(fname_v, dpi=200)
        plt.close(fig2)

# Correlaciones Spearman (heatmap) — todas las columnas mapeadas
# construir df de solo ítems numerados (ordenados por número)
ordered_nums = sorted([n for n in num_to_col.keys()])
ordered_cols = [num_to_col[n] for n in ordered_nums]
if ordered_cols:
    df_corr = df_num[ordered_cols].dropna()
    if not df_corr.empty:
        corr = df_corr.corr(method='spearman')
        fig, ax = plt.subplots(figsize=(10,8))
        cax = ax.imshow(corr.values, aspect='auto', interpolation='nearest')
        ax.set_xticks(np.arange(len(ordered_cols)))
        ax.set_yticks(np.arange(len(ordered_cols)))
        ax.set_xticklabels([str(n) for n in ordered_nums], rotation=90)
        ax.set_yticklabels([str(n) for n in ordered_nums])
        fig.colorbar(cax, ax=ax)
        ax.set_title("Matriz Spearman (ítems)")
        plt.tight_layout()
        fname_corr = os.path.join(GRAF_DIR, "matriz_correlaciones_spearman.png")
        fig.savefig(fname_corr, dpi=200)
        plt.close(fig)

print("  • Gráficos por ítem/dimensión generados en:", GRAF_DIR)

# ========== 6) SHAPIRO-WILK por dimensión ==========
print("\n[5] Prueba de normalidad (Shapiro-Wilk) por dimensión...")
rows_norm = []
for var in OPERACIONALIZACION:
    for dim, info in OPERACIONALIZACION[var].items():
        cols = [num_to_col[q] for q in info["items"] if q in num_to_col]
        if not cols:
            continue
        medianas_personas = df_num[cols].median(axis=1).dropna()
        n = medianas_personas.shape[0]
        if n < MIN_N_SHAPIRO:
            rows_norm.append({"Dimensión": dim, "N": n, "W": np.nan, "p_value": np.nan, "Normalidad": "Insuf. datos"})
            continue
        stat, pval = shapiro(medianas_personas)
        rows_norm.append({"Dimensión": dim, "N": n, "W": round(stat,4), "p_value": round(pval,4), "Normalidad": "NO (p<0.05)" if pval < 0.05 else "SÍ (p≥0.05)"})
df_norm = pd.DataFrame(rows_norm)
df_norm.to_csv(os.path.join(OUTPUT_FOLDER, "04_shapiro_normalidad.csv"), index=False, encoding='utf-8-sig')
print("  • CSV generado: 04_shapiro_normalidad.csv")

# ========== 7) WILCOXON ONE-SAMPLE (mediana vs 3) ==========
print("\n[6] Prueba Wilcoxon one-sample por dimensión (H0: mediana = 3, H1: mediana > 3)...")
rows_wil = []
for var in OPERACIONALIZACION:
    for dim, info in OPERACIONALIZACION[var].items():
        cols = [num_to_col[q] for q in info["items"] if q in num_to_col]
        if not cols:
            continue
        med_personas = df_num[cols].median(axis=1).dropna()
        n_total = med_personas.shape[0]
        if n_total < MIN_N_WILCOXON:
            rows_wil.append({"Variable": var, "Dimensión": dim, "N_total": n_total, "N_sin_empates": 0,
                             "Mediana_obs": np.nan, "IC95_inf": np.nan, "IC95_sup": np.nan, "W_stat": np.nan, "p_value": np.nan, "r": np.nan, "Decision": "Insuf. datos"})
            continue
        diferencias = med_personas - 3
        dif_no0 = diferencias[diferencias != 0]
        if dif_no0.shape[0] < MIN_N_WILCOXON:
            rows_wil.append({"Variable": var, "Dimensión": dim, "N_total": n_total, "N_sin_empates": dif_no0.shape[0],
                             "Mediana_obs": float(med_personas.median()), "IC95_inf": float(np.percentile(med_personas,2.5)), "IC95_sup": float(np.percentile(med_personas,97.5)),
                             "W_stat": np.nan, "p_value": np.nan, "r": np.nan, "Decision": "Insuf. datos (empates)"})
            continue
        try:
            stat, pval = wilcoxon(dif_no0, alternative='greater')
        except Exception as e:
            # fallback si no soporta 'alternative' (versiones antiguas de scipy)
            stat, pval_two = wilcoxon(dif_no0)
            pval = pval_two / 2.0
        # aproximación z para r (usada en varios trabajos)
        try:
            # isf para una cola
            z = stats.norm.isf(pval) if 0 < pval < 1 else 0.0
            r = abs(z) / np.sqrt(len(dif_no0))
        except:
            r = np.nan
        ic_inf = float(np.percentile(med_personas, 2.5))
        ic_sup = float(np.percentile(med_personas, 97.5))
        med_obs = float(med_personas.median())
        decision = "RECHAZA H0" if (pval < 0.05 and med_obs > 3) else "NO RECHAZA H0"
        rows_wil.append({
            "Variable": var,
            "Dimensión": dim,
            "N_total": int(n_total),
            "N_sin_empates": int(dif_no0.shape[0]),
            "Mediana_obs": round(med_obs, 3),
            "IC95_inf": round(ic_inf,3),
            "IC95_sup": round(ic_sup,3),
            "W_stat": int(stat) if not np.isnan(stat) else np.nan,
            "p_value": round(pval, 4) if not np.isnan(pval) else np.nan,
            "r": round(r, 3) if not np.isnan(r) else np.nan,
            "Decision": decision
        })

df_wil = pd.DataFrame(rows_wil)
df_wil.to_csv(os.path.join(OUTPUT_FOLDER, "05_wilcoxon_results.csv"), index=False, encoding='utf-8-sig')
print("  • CSV generado: 05_wilcoxon_results.csv")

# Gráfico resumen Wilcoxon (medianas + IC)
try:
    df_plot = df_wil.dropna(subset=["Mediana_obs"])
    if not df_plot.empty:
        fig, ax = plt.subplots(figsize=(10,6))
        y_pos = np.arange(len(df_plot))
        med_vals = df_plot["Mediana_obs"].values
        err_low = med_vals - df_plot["IC95_inf"].values
        err_up = df_plot["IC95_sup"].values - med_vals
        ax.barh(y_pos, med_vals, xerr=[err_low, err_up])
        ax.set_yticks(y_pos)
        ax.set_yticklabels([f"{r[:30]}" if isinstance(r,str) else r for r in df_plot["Dimensión"].values])
        ax.axvline(3, color='red', linestyle='--', linewidth=1)
        ax.set_xlabel("Mediana observada (IC 95%)")
        ax.set_title("Resumen Wilcoxon por dimensión")
        plt.tight_layout()
        fn = os.path.join(GRAF_DIR, "wilcoxon_resumen_medianas.png")
        fig.savefig(fn, dpi=200)
        plt.close(fig)
except Exception as e:
    print("  • Advertencia al generar gráfico de resumen Wilcoxon:", e)

# ========== 8) ITEM-TOTAL y ALFA SI SE ELIMINA ==========
print("\n[7] Análisis item-total y alfa si se elimina (por dimensión)...")
rows_item_total = []
rows_alpha_remove = []

for var in OPERACIONALIZACION:
    for dim, info in OPERACIONALIZACION[var].items():
        cols = [num_to_col[q] for q in info["items"] if q in num_to_col]
        if len(cols) < 2:
            continue
        df_dim = df_num[cols].dropna()
        if df_dim.empty:
            continue
        total_score = df_dim.sum(axis=1)
        for col in df_dim.columns:
            # item-total (Spearman)
            try:
                it_corr = df_dim[col].corr(total_score - df_dim[col], method='spearman')
            except:
                it_corr = np.nan
            rows_item_total.append({
                "Dimensión": dim,
                "Ítem_col": col,
                "Ítem_num": extraer_num_pregunta(col),
                "Item_Total_spearman": round(it_corr, 3) if not np.isnan(it_corr) else np.nan
            })
        # alfa si se elimina cada ítem
        for col in df_dim.columns:
            cols_minus = [c for c in df_dim.columns if c != col]
            a_minus = cronbach_alpha(df_dim[cols_minus]) if len(cols_minus) >= 2 else np.nan
            rows_alpha_remove.append({
                "Dimensión": dim,
                "Item_eliminado": extraer_num_pregunta(col),
                "Alpha_si_elimina": round(a_minus, 3) if not np.isnan(a_minus) else np.nan
            })

df_item_total = pd.DataFrame(rows_item_total)
df_item_total.to_csv(os.path.join(OUTPUT_FOLDER, "06_item_total_spearman.csv"), index=False, encoding='utf-8-sig')

df_alpha_remove = pd.DataFrame(rows_alpha_remove)
df_alpha_remove.to_csv(os.path.join(OUTPUT_FOLDER, "07_alpha_si_elimina.csv"), index=False, encoding='utf-8-sig')
print("  • CSVs generados: 06_item_total_spearman.csv, 07_alpha_si_elimina.csv")

# ========== 9) GENERAR REPORTE WORD resumen (opcional) ==========
print("\n[8] Generando reporte Word resumido (si python-docx funciona)...")
doc = Document()
doc.add_heading("REPORTE RESUMIDO - RESULTADOS TESIS", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.add_run("Diseño de un Sistema Web para la Gestión de Pagos - UNASAM\n").bold = True
p.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("1. Confiabilidad (resumen)", level=1)
doc.add_paragraph("Se adjunta CSV con Alfa Pearson y Alfa Ordinal por dimensión.")
try:
    doc.add_heading("2. Gráficos (seleccionados)", level=1)
    # insertar 2 gráficos representativos si existen
    fn1 = os.path.join(GRAF_DIR, "01_dimensiones_medianas.png")
    fn2 = os.path.join(GRAF_DIR, "wilcoxon_resumen_medianas.png")
    # si no existen, insertar los primeros dos archivos del directorio
    existing = sorted([os.path.join(GRAF_DIR, f) for f in os.listdir(GRAF_DIR) if f.lower().endswith(".png")])
    inserted = 0
    for fn in (fn1, fn2):
        if os.path.exists(fn):
            doc.add_picture(fn, width=Inches(6))
            inserted += 1
    if inserted == 0 and existing:
        doc.add_picture(existing[0], width=Inches(6))
        if len(existing) > 1:
            doc.add_picture(existing[1], width=Inches(6))
except Exception as e:
    print("  • Advertencia al insertar imágenes en Word:", e)

report_name = os.path.join(OUTPUT_FOLDER, "REPORTE_RESUMIDO_RESULTADOS.docx")
doc.save(report_name)
print("  • Reporte Word generado:", report_name)

# ========== 10) RESUMEN EN CONSOLA ==========
print("\nAnálisis finalizado. Archivos generados en:", OUTPUT_FOLDER)
print("Principales CSV generados:")
for f in sorted(os.listdir(OUTPUT_FOLDER)):
    if f.lower().endswith(".csv") or f.lower().endswith(".docx"):
        print("  -", f)
print("\nGráficos en:", GRAF_DIR)
print("  - Ej: item_1_frecuencia.png, dim_*_violin.png, matriz_correlaciones_spearman.png, wilcoxon_resumen_medianas.png")

print("\nFIN del script. Revisa los CSVs y sube aquí si quieres que redacte o integre las tablas/figuras en tu Capítulo IV.")
